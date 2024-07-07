#!/usr/bin/env python

# import the necessary packages
import numpy as np
import cv2
import sys
from docx import Document
from docx.shared import Cm

## INPUT PARAMETERS ##

ID=0
width_in_cm = 19
height_in_cm = 19
Marker_name = "id_"+str(ID)+"_aruco_marker.png"

## INPUT PARAMETERS ##


doc = Document()
section = doc.sections[0]
section.left_margin = Cm(1)
section.right_margin = Cm(1)
image_path=Marker_name


ARUCO_DICT = {"dict": cv2.aruco.DICT_ARUCO_MIP_36h12}
arucoDict = cv2.aruco.getPredefinedDictionary(ARUCO_DICT["dict"])
tag = np.zeros((300, 300, 1), dtype="uint8")
cv2.aruco.generateImageMarker(arucoDict, ID, 300, tag, 1)
cv2.imwrite(Marker_name, tag)
doc.add_picture(image_path, width=Cm(width_in_cm), height=Cm(height_in_cm))
doc.save(Marker_name+".docx")


#Loop Version#

#for i in range(16):
#    ID=i
#    Marker_name = "id_"+str(ID)+"_aruco_marker.png"
#
#    arucoDict = cv2.aruco.getPredefinedDictionary(ARUCO_DICT["dict"])
#
#    tag = np.zeros((300, 300, 1), dtype="uint8")
#    cv2.aruco.generateImageMarker(arucoDict, ID, 300, tag, 1)
#    cv2.imwrite(Marker_name, tag)
#    image_path=Marker_name
#    doc.add_picture(image_path, width=Cm(width_in_cm), height=Cm(height_in_cm))
#doc.save("aruco_marker.docx")

	
